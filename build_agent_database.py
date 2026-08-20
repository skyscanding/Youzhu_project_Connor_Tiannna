"""Build a searchable SQLite database for the city-renewal research agent.

The source project is treated as read-only.  This script never modifies it;
it creates a new database with canonical documents, section-level passages,
crawler records, provenance, and SQLite FTS5 indexes.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sqlite3
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path


SOURCE_DB = "cityrenewal.db"


def normalise_title(value: str) -> str:
    return re.sub(r"[\s\u3000（）()【】\[\]《》,.，:：;；'\"“”‘’\-—–]", "", value or "").lower()


def slug(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:20]


def read_text(root: Path, relative_path: str) -> str:
    if not relative_path:
        return ""
    file_path = root / relative_path
    if not file_path.is_file():
        return ""
    return file_path.read_text(encoding="utf-8", errors="replace")


def read_docx_text(path: Path) -> str:
    """Extract paragraphs from a user-provided DOCX and preserve simple headings."""
    from docx import Document

    document = Document(path)
    rendered: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if re.match(r"^[一二三四五六七八九十]+、", text):
            rendered.append("# " + text)
        elif re.match(r"^（[一二三四五六七八九十]+）", text):
            rendered.append("## " + text)
        else:
            rendered.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rendered.append(" | ".join(cells))
    return "\n\n".join(rendered)


def pick(values: list[str]) -> str:
    """Prefer the most informative non-empty value."""
    values = [value.strip() for value in values if value and value.strip()]
    return max(values, key=len) if values else ""


def classify_jurisdiction(city: str, title: str, category: str, authority: str) -> tuple[str, str]:
    text = " ".join((city, title, category, authority))
    if any(token in text for token in ("国家", "中华人民共和国", "住房和城乡建设部", "自然资源部", "财政部")):
        return "national", "全国"
    if "北京市" in city and "区" in city:
        return "district", city
    if "北京" in text:
        district = re.search(r"(?:北京市)?([\u4e00-\u9fff]{2,5}区)", city or title)
        return ("district", f"北京市{district.group(1)}") if district else ("municipal", "北京市")
    for city_name in ("深圳市", "上海市", "广州市", "重庆市", "天津市", "南京市", "杭州市", "苏州市"):
        if city_name[:2] in text:
            return "other_local", city_name
    if city:
        return "other_local", city
    return "unknown", "未标注"


def supplemental_metadata(root: Path) -> dict[str, list[dict[str, str]]]:
    """Read the two human-curated indexes to recover metadata absent from v1."""
    result: dict[str, list[dict[str, str]]] = defaultdict(list)
    for filename in ("国家级法律规范.json", "手动检索结果.json"):
        path = root / filename
        if not path.is_file():
            continue
        payload = json.loads(path.read_text(encoding="utf-8"))
        record_lists = [value for value in payload.values() if isinstance(value, list) and value and isinstance(value[0], dict)]
        for records in record_lists:
            for record in records:
                title = record.get("标题", "")
                if not title:
                    continue
                result[normalise_title(title)].append({
                    "authority": record.get("信源名称", ""),
                    "document_number": record.get("发文号", ""),
                    "published_date": record.get("日期", ""),
                    "original_path": record.get("原文", ""),
                    "official_url": record.get("记录URL", ""),
                    "level": record.get("层级", ""),
                    "source_label": record.get("数据来源", ""),
                })
    return result


def load_supplemental_text(path: Path | None) -> dict[str, dict[str, str]]:
    if path is None:
        return {}
    records = json.loads(path.read_text(encoding="utf-8"))
    return {normalise_title(record["target_title"]): record for record in records}


def classify_document(title: str, category: str, source_type: str) -> tuple[str, str]:
    text = " ".join((title, category, source_type))
    if any(token in text for token in ("法", "条例")) and "办法" not in text:
        kind = "law_or_regulation"
    elif "规范" in text or "标准" in text:
        kind = "technical_standard"
    elif any(token in text for token in ("指南", "细则", "办法", "意见", "方案", "计划")):
        kind = "policy_or_guideline"
    elif any(token in text for token in ("新闻", "访谈", "报道")):
        kind = "official_news_or_case"
    else:
        kind = "document"
    if kind == "official_news_or_case":
        return kind, "reference"
    return kind, "direct_basis" if source_type != "手动检索" else "needs_verification"


def split_sections(text: str) -> list[tuple[str, int, str]]:
    """Split Markdown into answerable passages while retaining heading context."""
    lines = text.splitlines()
    sections: list[tuple[str, int, str]] = []
    heading_stack: list[str] = []
    body: list[str] = []
    level = 0

    def flush() -> None:
        content = "\n".join(body).strip()
        if content:
            sections.append((" > ".join(heading_stack), level, content))

    for line in lines:
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            flush()
            body = []
            level = len(match.group(1))
            heading_stack = heading_stack[: level - 1] + [match.group(2)]
        else:
            body.append(line)
    flush()
    return sections or [("", 0, text.strip())] if text.strip() else []


SCHEMA = """
PRAGMA foreign_keys = ON;
CREATE TABLE metadata (key TEXT PRIMARY KEY, value TEXT NOT NULL);
CREATE TABLE documents (
  id TEXT PRIMARY KEY,
  title TEXT NOT NULL,
  title_normalized TEXT NOT NULL UNIQUE,
  document_type TEXT NOT NULL,
  evidence_tier TEXT NOT NULL CHECK (evidence_tier IN ('direct_basis','reference','needs_verification')),
  jurisdiction_level TEXT NOT NULL,
  jurisdiction_name TEXT NOT NULL,
  issuing_authority TEXT NOT NULL DEFAULT '',
  document_number TEXT NOT NULL DEFAULT '',
  published_date TEXT NOT NULL DEFAULT '',
  effective_date TEXT NOT NULL DEFAULT '',
  validity_status TEXT NOT NULL DEFAULT 'unknown',
  category TEXT NOT NULL DEFAULT '',
  domain TEXT NOT NULL DEFAULT '',
  city_update_relevance TEXT NOT NULL DEFAULT '',
  abstract TEXT NOT NULL DEFAULT '',
  content TEXT NOT NULL DEFAULT '',
  original_path TEXT NOT NULL DEFAULT '',
  official_url TEXT NOT NULL DEFAULT '',
  content_hash TEXT NOT NULL,
  created_at TEXT NOT NULL
);
CREATE TABLE document_provenance (
  id INTEGER PRIMARY KEY,
  document_id TEXT NOT NULL REFERENCES documents(id),
  original_id TEXT NOT NULL,
  source_type TEXT NOT NULL DEFAULT '',
  source_label TEXT NOT NULL DEFAULT '',
  original_path TEXT NOT NULL DEFAULT '',
  imported_at TEXT NOT NULL,
  UNIQUE(document_id, original_id)
);
CREATE TABLE document_sections (
  id INTEGER PRIMARY KEY,
  document_id TEXT NOT NULL REFERENCES documents(id),
  ordinal INTEGER NOT NULL,
  heading_path TEXT NOT NULL DEFAULT '',
  heading_level INTEGER NOT NULL DEFAULT 0,
  content TEXT NOT NULL,
  UNIQUE(document_id, ordinal)
);
CREATE TABLE crawler_sources (
  id TEXT PRIMARY KEY,
  name TEXT NOT NULL DEFAULT '',
  jurisdiction_level TEXT NOT NULL DEFAULT '',
  source_url TEXT NOT NULL DEFAULT '',
  crawl_ok INTEGER NOT NULL DEFAULT 0
);
CREATE TABLE web_records (
  id INTEGER PRIMARY KEY,
  source_id TEXT NOT NULL REFERENCES crawler_sources(id),
  title TEXT NOT NULL,
  record_type TEXT NOT NULL DEFAULT 'crawl_lead',
  jurisdiction_level TEXT NOT NULL DEFAULT '',
  jurisdiction_name TEXT NOT NULL DEFAULT '',
  evidence_tier TEXT NOT NULL DEFAULT 'needs_verification',
  url TEXT NOT NULL DEFAULT '',
  published_date TEXT NOT NULL DEFAULT '',
  document_number TEXT NOT NULL DEFAULT '',
  first_seen TEXT NOT NULL DEFAULT '',
  last_seen TEXT NOT NULL DEFAULT '',
  content TEXT NOT NULL DEFAULT '',
  fetch_status TEXT NOT NULL DEFAULT '',
  fetch_date TEXT NOT NULL DEFAULT '',
  content_hash TEXT NOT NULL,
  UNIQUE(source_id, url, title)
);
CREATE INDEX idx_documents_scope ON documents(jurisdiction_level, jurisdiction_name, evidence_tier);
CREATE INDEX idx_documents_date ON documents(published_date);
CREATE INDEX idx_sections_document ON document_sections(document_id);
CREATE INDEX idx_web_scope ON web_records(jurisdiction_level, jurisdiction_name, published_date);
-- Trigram indexing supports unsegmented Chinese text without a third-party tokenizer.
CREATE VIRTUAL TABLE documents_fts USING fts5(title, abstract, content, jurisdiction_name, evidence_tier, content='documents', content_rowid='rowid', tokenize='trigram');
CREATE VIRTUAL TABLE sections_fts USING fts5(heading_path, content, content='document_sections', content_rowid='id', tokenize='trigram');
CREATE VIRTUAL TABLE web_records_fts USING fts5(title, content, jurisdiction_name, evidence_tier, content='web_records', content_rowid='id', tokenize='trigram');
CREATE TRIGGER documents_ai AFTER INSERT ON documents BEGIN
  INSERT INTO documents_fts(rowid,title,abstract,content,jurisdiction_name,evidence_tier) VALUES (new.rowid,new.title,new.abstract,new.content,new.jurisdiction_name,new.evidence_tier);
END;
CREATE TRIGGER documents_ad AFTER DELETE ON documents BEGIN
  INSERT INTO documents_fts(documents_fts,rowid,title,abstract,content,jurisdiction_name,evidence_tier) VALUES ('delete',old.rowid,old.title,old.abstract,old.content,old.jurisdiction_name,old.evidence_tier);
END;
CREATE TRIGGER documents_au AFTER UPDATE ON documents BEGIN
  INSERT INTO documents_fts(documents_fts,rowid,title,abstract,content,jurisdiction_name,evidence_tier) VALUES ('delete',old.rowid,old.title,old.abstract,old.content,old.jurisdiction_name,old.evidence_tier);
  INSERT INTO documents_fts(rowid,title,abstract,content,jurisdiction_name,evidence_tier) VALUES (new.rowid,new.title,new.abstract,new.content,new.jurisdiction_name,new.evidence_tier);
END;
CREATE TRIGGER sections_ai AFTER INSERT ON document_sections BEGIN
  INSERT INTO sections_fts(rowid,heading_path,content) VALUES(new.id,new.heading_path,new.content);
END;
CREATE TRIGGER web_records_ai AFTER INSERT ON web_records BEGIN
  INSERT INTO web_records_fts(rowid,title,content,jurisdiction_name,evidence_tier) VALUES(new.id,new.title,new.content,new.jurisdiction_name,new.evidence_tier);
END;
"""


def build(source_root: Path, output: Path, supplemental_docx_paths: list[Path] | None = None, supplemental_text_path: Path | None = None) -> None:
    source_db = source_root / SOURCE_DB
    if not source_db.is_file():
        raise FileNotFoundError(f"Source database not found: {source_db}")
    if output.exists():
        output.unlink()
    source = sqlite3.connect(f"file:{source_db.as_posix()}?mode=ro", uri=True)
    source.row_factory = sqlite3.Row
    target = sqlite3.connect(output)
    target.executescript(SCHEMA)
    now = datetime.now(timezone.utc).isoformat()
    supplemental = supplemental_metadata(source_root)
    text_records = load_supplemental_text(supplemental_text_path)

    rows = list(source.execute("SELECT * FROM documents ORDER BY id"))
    grouped: dict[str, list[sqlite3.Row]] = defaultdict(list)
    for row in rows:
        grouped[normalise_title(row["title"])].append(row)

    supplemental_docx: dict[str, tuple[Path, str]] = {}
    for path in supplemental_docx_paths or []:
        if not path.is_file():
            raise FileNotFoundError(f"Supplemental DOCX not found: {path}")
        text = read_docx_text(path)
        matches = [title_key for title_key in grouped if title_key and title_key in normalise_title(text)]
        if len(matches) != 1:
            raise ValueError(f"Could not uniquely match supplemental DOCX to one database title: {path}")
        supplemental_docx[matches[0]] = (path, text)

    for title_key, group in grouped.items():
        extra = supplemental.get(title_key, [])
        text_record = text_records.get(title_key)
        title = pick([row["title"] for row in group])
        content_by_row = {row["id"]: read_text(source_root, row["md_file"]) for row in group}
        for index, record in enumerate(extra):
            extra_path = record["original_path"]
            extra_content = read_text(source_root, extra_path)
            if len(extra_content) > len(content_by_row.get("__extra__", "")):
                content_by_row["__extra__"] = extra_content
        docx_entry = supplemental_docx.get(title_key)
        if docx_entry:
            content_by_row["__supplemental_docx__"] = docx_entry[1]
        if text_record:
            content_by_row["__supplemental_text__"] = text_record["content"]
        best_content_id = max(content_by_row, key=lambda item: len(content_by_row[item]))
        content = content_by_row[best_content_id]
        category = pick([row["cat"] for row in group])
        authority = text_record["source_label"] if text_record else pick([row["pub_org"] for row in group] + [record["authority"] for record in extra])
        city = pick([row["city"] for row in group])
        level, jurisdiction = classify_jurisdiction(city, title, category, authority)
        extra_levels = [record["level"] for record in extra]
        if any("国家" in item for item in extra_levels):
            level, jurisdiction = "national", "全国"
        source_types = [row["source_type"] for row in group]
        kind, tier = classify_document(title, category, pick(source_types))
        if text_record:
            kind, tier = text_record["document_type"], text_record["evidence_tier"]
        canonical_id = f"doc-{slug(title_key)}"
        target.execute(
            """INSERT INTO documents VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (canonical_id, title, title_key, kind, tier, level, jurisdiction, authority,
             text_record.get("document_number", "") if text_record else pick([row["docno"] for row in group] + [record["document_number"] for record in extra]), text_record["published_date"] if text_record else pick([row["pub_date"] for row in group] + [record["published_date"] for record in extra]),
             text_record.get("effective_date", "") if text_record else pick([row["impl_date"] for row in group]), text_record.get("validity_status", "unknown") if text_record else "unknown", category,
             pick([row["domain"] for row in group]), pick([row["city_update"] for row in group]),
             pick([row["abstract"] for row in group]), content,
             str(supplemental_text_path) if text_record and supplemental_text_path else pick([row["md_file"] for row in group] + [record["original_path"] for record in extra] + ([str(docx_entry[0])] if docx_entry else [])),
             text_record.get("official_url", "") if text_record else pick([record["official_url"] for record in extra]), hashlib.sha256(content.encode("utf-8")).hexdigest(), now),
        )
        for row in group:
            target.execute(
                "INSERT INTO document_provenance(document_id,original_id,source_type,source_label,original_path,imported_at) VALUES (?,?,?,?,?,?)",
                (canonical_id, row["id"], row["source_type"], row["source"], row["md_file"], now),
            )
        for index, record in enumerate(extra):
            target.execute(
                "INSERT OR IGNORE INTO document_provenance(document_id,original_id,source_type,source_label,original_path,imported_at) VALUES (?,?,?,?,?,?)",
                (canonical_id, f"curated-{index}-{slug(title_key)}", "curated_index", record["source_label"], record["original_path"], now),
            )
        if docx_entry:
            target.execute(
                "INSERT OR IGNORE INTO document_provenance(document_id,original_id,source_type,source_label,original_path,imported_at) VALUES (?,?,?,?,?,?)",
                (canonical_id, f"docx-{slug(str(docx_entry[0]))}", "user_provided_docx", "用户补充原始 DOCX", str(docx_entry[0]), now),
            )
        if text_record and supplemental_text_path:
            target.execute(
                "INSERT OR IGNORE INTO document_provenance(document_id,original_id,source_type,source_label,original_path,imported_at) VALUES (?,?,?,?,?,?)",
                (canonical_id, f"text-{slug(title_key)}", "user_provided_text", text_record["source_label"], str(supplemental_text_path), now),
            )
        for ordinal, (heading, heading_level, section_content) in enumerate(split_sections(content), 1):
            target.execute(
                "INSERT INTO document_sections(document_id,ordinal,heading_path,heading_level,content) VALUES (?,?,?,?,?)",
                (canonical_id, ordinal, heading, heading_level, section_content),
            )

    for row in source.execute("SELECT * FROM sources ORDER BY id"):
        target.execute(
            "INSERT INTO crawler_sources(id,name,jurisdiction_level,source_url,crawl_ok) VALUES (?,?,?,?,?)",
            (row["id"], row["name"], row["level"], row["url"], row["ok"]),
        )
    source_map = {row["id"]: row for row in source.execute("SELECT * FROM sources")}
    for row in source.execute("SELECT * FROM snapshots ORDER BY id"):
        source_row = source_map[row["source_id"]]
        source_name = source_row["name"]
        level, jurisdiction = classify_jurisdiction("", source_name, "", source_name)
        if row["source_id"].startswith("bj-"):
            district = re.search(r"([\u4e00-\u9fff]{2,5}区)", source_name)
            level, jurisdiction = ("district", f"北京市{district.group(1)}") if district else ("municipal", "北京市")
        content = row["content"] or ""
        title = row["title"]
        record_type = "case_or_news" if any(token in title for token in ("访谈", "报道", "改造", "活动")) else "crawl_lead"
        target.execute(
            """INSERT OR IGNORE INTO web_records(source_id,title,record_type,jurisdiction_level,jurisdiction_name,evidence_tier,url,published_date,document_number,first_seen,last_seen,content,fetch_status,fetch_date,content_hash)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (row["source_id"], title, record_type, level, jurisdiction, "needs_verification", row["url"], row["date"],
             row["doc_number"], row["first_seen"], row["last_seen"], content, row["fetch_status"], row["fetch_date"],
             hashlib.sha256((title + content).encode("utf-8")).hexdigest()),
        )

    metadata = {
        "schema_version": "2.0",
        "purpose": "北京市城市更新政策与案例研究 Agent；不输出自动合规结论",
        "source_database": str(source_db),
        "built_at": now,
        "evidence_tiers": json.dumps({
            "direct_basis": "可作为直接依据的法规、规范、政策或指南；仍需核验效力状态与原文。",
            "reference": "案例、新闻或其他地区实践，仅作研究参考。",
            "needs_verification": "爬虫或人工导入线索，使用前需核验原文、效力和适用范围。",
        }, ensure_ascii=False),
        "search_order": "全国 > 北京市 > 用户指定区 > 北京其他区参考 > 外地参考",
    }
    target.executemany("INSERT INTO metadata(key,value) VALUES (?,?)", metadata.items())
    target.commit()
    target.execute("PRAGMA optimize")
    target.close()
    source.close()


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-root", type=Path, required=True, help="Read-only project containing cityrenewal.db")
    parser.add_argument("--output", type=Path, required=True, help="New SQLite v2 database path")
    parser.add_argument("--supplemental-docx", type=Path, action="append", default=[], help="DOCX whose text should be matched and imported into an existing document")
    parser.add_argument("--supplemental-text", type=Path, help="JSON list of user-supplied texts matched to existing document titles")
    args = parser.parse_args()
    build(args.source_root.resolve(), args.output.resolve(), [path.resolve() for path in args.supplemental_docx], args.supplemental_text.resolve() if args.supplemental_text else None)
    print(args.output.resolve())


if __name__ == "__main__":
    main()

